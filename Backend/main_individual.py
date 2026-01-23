#This Third version is created for emergent solution new entity

import random
import re

import cv2
import googlemaps
# load_qa_chain has been removed in newer versions of langchain
# from langchain.chains import load_qa_chain
# try:
#     from langchain.chat_models import ChatOpenAI
# except ImportError:
#     from langchain_community.chat_models import ChatOpenAI
# from langchain_community.document_loaders import CSVLoader, DirectoryLoader
# try:
#     from langchain.llms import AzureOpenAI
# except ImportError:
#     from langchain_community.llms import AzureOpenAI

from skimage import filters
from transformers import pipeline
from sentence_transformers import SentenceTransformer, util
from pyzbar.pyzbar import decode
import xml.etree.ElementTree as ET
import pandas as pd
import numpy as np
import compare_faces
from flask import Flask, request,send_file,jsonify
from flask_cors import CORS
import matplotlib.pyplot as plt
import seaborn as sns
from reportlab.pdfgen import canvas
import base64
import io
import openai
import spacy
from googlesearch import search
from bs4 import BeautifulSoup
from urllib.request import urlopen
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import Spacer
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.platypus.flowables import KeepInFrame
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# ...existing code...
from passporteye import read_mrz
from datetime import datetime
import docx

# Load environment variables from the .env file
import os
from dotenv import load_dotenv
load_dotenv()

# load_summarize_chain has been deprecated in newer langchain versions
# try:
#     from langchain.chains.summarize import load_summarize_chain
# except ImportError:
#     from langchain_community.chains.summarize import load_summarize_chain


# try:
#     from langchain.text_splitter import CharacterTextSplitter
# except ImportError:
#     from langchain_text_splitters import CharacterTextSplitter
# try:
#     from langchain.text_splitter import RecursiveCharacterTextSplitter
# except ImportError:
#     from langchain_text_splitters import RecursiveCharacterTextSplitter
# try:
#     from langchain.prompts import PromptTemplate
# except ImportError:
#     from langchain_core.prompts import PromptTemplate

from pdfminer.high_level import extract_text
from io import BytesIO
import mysql.connector
from fuzzywuzzy import fuzz
from flask_mail import Mail, Message
import os
import ftfy
import skimage.filters as filters
from email_send import email_send
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
# LLMChain and StuffDocumentsChain have been moved/deprecated in newer langchain versions
# try:
#     from langchain.chains.llm import LLMChain
# except ImportError:
#     from langchain_community.chains.llm import LLMChain
# try:
#     from langchain.chains.combine_documents.stuff import StuffDocumentsChain
# except ImportError:
#     from langchain_community.chains.combine_documents.stuff import StuffDocumentsChain
import time


os.environ['OPENAI_API_KEY'] = os.getenv("OPENAI_API_KEY")
os.environ['CURL_CA_BUNDLE'] = ''
os.environ['REQUESTS_CA_BUNDLE'] = ''
key = os.getenv("GOOGLE_API_KEY")
os.environ["GOOGLE_API_KEY"] = key
genai.configure(api_key = key)

app = Flask(__name__)
CORS(app)

app.config['MAIL_SERVER']='smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USERNAME'] = os.getenv("MAIL_USERNAME")
app.config['MAIL_PASSWORD'] = os.getenv("MAIL_PASSWORD")
app.config['MAIL_USE_TLS'] = False
app.config['MAIL_USE_SSL'] = True
mail = Mail(app)


class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata

def dbConnection():
    mydb = mysql.connector.connect(
        host="localhost",
        user="root",
        password="Mysqlroot"  ## Update the password as per your local mysql configuration
    )
    print(mydb)
    cursor = mydb.cursor()
    return mydb,cursor

def image_decoding(image):
    image_data = base64.b64decode(image)
    im_arr = np.frombuffer(image_data, dtype=np.uint8)  # im_arr is one-dim Numpy array
    print(im_arr.shape)
    decoded_image = cv2.imdecode(im_arr, flags=cv2.IMREAD_COLOR)
    return decoded_image

def extract_aadhar_details(code):
    try:
        qrData = code[0].data
        print(qrData, type(qrData))
    except:
        return "Please upoload clear image"
    xml_string = qrData.decode('utf-8')
    # Parse the XML string
    root = ET.fromstring(xml_string)
    # Extract attributes
    attributes = root.attrib
    print(f'attributes:{attributes}')
    address_keys = []
    between_keys = False
    for key in attributes:
        if key == 'co':
            between_keys = True
        elif key == 'dob':
            break
        if between_keys:
            address_keys.append(key)
    # Concatenate the values of address keys
    address_values = ','.join(str(attributes.get(key, '')) for key in address_keys[1:])
    if attributes.get('gender') == 'M':
        attributes['gender'] = 'Male'
    elif attributes.get('gender') == 'F':
        attributes['gender'] = 'Female'
    df = pd.read_excel('Aadhar_Matching_data.xlsx')
    aadhar_li = df['AadharNumber'].to_list()
    aadhar_comparison = "Not Matched"
    for i in aadhar_li:
        if str(attributes.get('uid'))==str(i):
            aadhar_comparison = "Matched"
            break
    details = {'Aadhar_number': attributes.get('uid'), 'Name': attributes['name'], 'Gender': attributes['gender'],
               'CO': attributes['co'], 'DOB': attributes['dob'], 'Address':address_values, 'District':attributes['dist'], 'State':
    attributes['state'],'aadhar_comparison':aadhar_comparison,'name_screening':"Matched - 2 Hits","adverse_media_screening":"Matched - Negative Analysis-80%,Positive Analysis–18%,Neutral Analysis-2%","transaction_monitoring":'Matched'}

    return details

def create_pdf_with_list(output_filename,heading,item_list):
    c = canvas.Canvas(output_filename, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(250, 750, heading)
    c.setFont("Helvetica", 12)
    # Define starting position for writing
    x, y = 50, 700
    wrap_width = 500
    # Set up styles for paragraphs
    styles = getSampleStyleSheet()
    style = styles["Normal"]

    # Write the list items with word wrapping
    for item in item_list:
        paragraph = Paragraph(f"- {item}", style)
        paragraph.wrap(wrap_width, 600)  # Limit the height to prevent overlapping
        paragraph.drawOn(c, x, y)
        y -= paragraph.height + 10
    c.save()

def fetch_and_scrape(url,keyword):
    try:
        search_kwd = keyword
        page = urlopen(url)
        html = page.read().decode("utf-8")
        soup = BeautifulSoup(html, "html.parser")
        p_text = soup.select('p')
        main_text = ''
        for line in p_text:
            main_text = main_text + line.text
        #print(f'main_content:{main_text}')

        # text_splitter = CharacterTextSplitter(chunk_overlap=0, chunk_size=10)
        # chunk_text = text_splitter.split_text(main_text)
        chunk_text = [main_text[i:i + 3000] for i in range(0, len(main_text), 3000)]
        #print(len(chunk_text),type(chunk_text))
        relevant_sentences = []
        for i in chunk_text:
            relevant_sentences.extend(filter_sentences_with_gpt3(i,search_kwd))
        cleaned_content = '\n'.join(relevant_sentences)
        #print(f'cleaned_content:{cleaned_content}')
        return cleaned_content


    except Exception as e:
        print(f"Error while fetching/scraping {url}: {str(e)}")
        return None

def filter_sentences_with_gpt3(content,search_query):
    try:
        openai.api_key = os.getenv("OPENAI_API_KEY1")
        openai.api_base = "https://bigaidea.openai.azure.com/"
        openai.api_type = 'azure'
        openai.api_version = '2022-12-01'
        engine = 'deployment-8c300b6d6ee747668b102e938fe54f70'
        prompt = f"Filter the relevant sentences from the following content: \n\n{content}\n\nSearch Query: {search_query}\n\nRelevant Sentences:\n"
        response = openai.Completion.create(
            model="text-davinci-003",
            prompt=prompt,
            max_tokens=500,
            temperature=0,
            engine=engine
        )
        # Extract and return the relevant sentences
        relevant_sentences = response.choices[0].text.strip().split('\n')
        return relevant_sentences
    except Exception as e:
        print(f"Error while filtering sentences with GPT-3: {str(e)}")
        return []

def sentence_parsing(filename):
    nlp = spacy.load("en_core_web_sm")
    temp  = open(filename,'r',encoding="utf-8")
    main_text = ''
    for i in temp:
        if not (i.startswith('URL') or i.startswith('---') or i.startswith('The site is') or i.startswith('Content:')):
            main_text = main_text+i
    print(main_text)
    doc = nlp(main_text)
    li = []
    for sent in doc.sents:
        print(sent.text)
        print('***************************')
        li.append(sent.text)
    print(li)
    df = pd.DataFrame({'FreeText':li})
    df.to_excel('SentimentalAnlysisInputData_afterTokenization.xlsx',index=False)
    print('Input Excel Created for SentimentalAnlysis!!')

def name_screening_model(input_data,screen_list):
    model = SentenceTransformer('shahrukhx01/paraphrase-mpnet-base-v2-fuzzy-matcher')

    sentence_embeddings1 = model.encode(input_data, convert_to_tensor=True)
    sentence_embeddings2 = model.encode(screen_list, convert_to_tensor=True)
    threshold_value = 0.75
    result = dict()
    values = []
    for i, embedding1 in enumerate(sentence_embeddings1):
        for j, embedding2 in enumerate(sentence_embeddings2):
            temp_dict = dict()
            cosine_score = util.cos_sim(embedding1, embedding2)
            print(cosine_score.item())
            if cosine_score.item() >= threshold_value:
                temp_dict['Name'] = screen_list[j].split(',')[0]
                temp_dict['Address'] = ','.join(screen_list[j].split(',')[1:])
                temp_dict['Score'] = round(cosine_score.item(), 2)
                values.append(temp_dict)

    print(values)
    result['values'] = values
    result['No_of_hits'] = len(values)
    print(result)
    return result


@app.route('/upload_aadhar', methods=['POST'])
def upload_aadhar():
    data = request.get_json()
    aadhar_image = image_decoding(data['aadhar'])
    gray_image = cv2.cvtColor(aadhar_image, cv2.COLOR_BGR2GRAY)
    code = decode(gray_image)
    details = extract_aadhar_details(code)
    print(type(details),details)
    return details

@app.route('/upload_profile', methods=['POST'])
def upload_profile():
    data = request.get_json()
    aadhar_image = image_decoding(data['aadhar'])
    recent_image = image_decoding(data['profile'])

    score_val, msg_val, type_val = compare_faces.run(aadhar_image, recent_image)
    print(score_val, msg_val, type_val)
    response = dict()
    response['Comparison'] = msg_val
    #response['Score'] = score_val
    return response

'''
@app.route('/qna', methods=['POST'])
def qna():
    req = request.get_json()
    question = req['question']
    if req['tab'] == 'individual':
    # Load your CSV or JSON data into a string
        with open("demo_data_for_qna_and_Report_summary.csv", "r") as file:
            data = file.read()
    else:
        #llm = ChatOpenAI(temperature=0.0, model_name='gpt-3.5-turbo')
        llm = ChatGoogleGenerativeAI(model="gemini-pro", 
                                     convert_system_message_to_human=True,
                                     temperature=0)
        file_path = 'entity_emergent_qna_input_data.pdf'
        text = extract_text(file_path)
        #print(text)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        splits = text_splitter.split_text(text)
        docs = text_splitter.create_documents(splits)
        # template = f"""Use the following pieces of context to answer the question at the end.
        #         Ellaborate the answer breifly and if answer is not there try to answer it in gracefully as Information not available
        #         If asked about UBO give name ,along with its percentage of share
        #         If asked about subsidiaries provide only the name of subsidiaries and its country who are operating in Jordan,Turkey,Uganda,United Arab Emirates,Yemen
        #         Question:{question}"""
        template = f"""Use the text below to answer the question. \n
                       Question:{question}"""
        prompt_template = template + "{text}"
        prompt = PromptTemplate.from_template(prompt_template)
        llm_chain = LLMChain(llm=llm, prompt=prompt)
        stuff_chain = StuffDocumentsChain(llm_chain=llm_chain, document_variable_name="text")
        try:
            answer =stuff_chain.run(docs)
        except Exception as e:
            time.sleep(120)
            answer =stuff_chain.run(docs)
                
        return {"Answer":answer}
'''

@app.route('/download_report', methods=['GET'])
def download_report():
    df = pd.read_excel('demo_data_for_qna_and_Report_summary.xlsx')
    categorical_columns = ['Type', 'Gender', 'State', 'District']

    chart_filenames = []

    for col in categorical_columns:
        plt.figure(figsize=(20, 10), dpi=200)  # Adjust figsize and dpi
        sns.countplot(data=df, x=col)
        plt.title(f'{col}Distribution')
        plt.xlabel(col)
        plt.ylabel('Count')
        chart_filename = f'{col}_chart.png'
        plt.savefig(chart_filename)
        plt.close()
        chart_filenames.append((chart_filename, col))
    #to create hue charts with Type
    for col in categorical_columns[1:]:
        plt.figure(figsize=(20, 10), dpi=200)  # Adjust figsize and dpi
        sns.countplot(data=df, x=col, hue='Type')
        plt.title(f'{col} with Type Distribution')
        plt.xlabel(col)
        plt.ylabel('Count')
        chart_filename = f'{col}_hue_chart.png'
        plt.savefig(chart_filename)
        plt.close()
        chart_filenames.append((chart_filename, col))

    # Generate PDF report
    c = canvas.Canvas('chart_report_new.pdf', pagesize=letter)
    # c.drawCentredString(300, 750, "Report")  # Centered title
    # c.showPage()

    for chart_filename, col in chart_filenames:
        # Start a new page for each chart
        c.drawCentredString(300, 750, f"{col} Distribution")  # Centered chart title
        c.drawImage(chart_filename, 50, 300, width=500, height=300)  # Adjust positioning
        c.showPage()
    c.save()
    print('Report PDF created!')
    return send_file('chart_report_new.pdf',mimetype='application/pdf')

@app.route('/adverse_media_scraper', methods=['POST'])
def adverse_media_scraper():
    req = request.get_json()
    if 'person' in req:
        #keyword = req['person'] + "'s "  + req['keyword']
        keyword = "Sadam Anwar" + "'s " + req['keyword']
    else:
        keyword = req['keyword'] # Keyword has to form here
    try:
        search_results = search(keyword,num_results=2, lang="en")
        filename = f'scraper_output_{str(random.randint(1,100))}.txt'# Delete the exsisting file for every run
        print(filename)
        with open(filename, "a", encoding="utf-8") as file:
            for index, url in enumerate(search_results, start=1):
                print(f"Scraping content from URL {index}: {url}")
                file.write(f"URL {index}: {url}\n")
                content = fetch_and_scrape(url,keyword)
                if content:
                    # print(f'Content:{content}')
                    # print("\n" + "-" * 50 + "\n")
                    file.write("Content:\n")
                    file.write(content)
                    file.write("\n" + "-" * 50 + "\n")
                else:
                    file.write("The site is Forbidden to access the content!\n")
                    file.write("\n" + "-" * 50 + "\n")
        sentence_parsing(filename)
        return send_file(filename, as_attachment=True)

    except Exception as e:
        print(f"Error during the search and scrape process: {str(e)}")
        return "Please check the internet connection...Try once when its stable!!"

'''
@app.route('/adverse_media_summarization', methods=['POST'])
def adverse_media_summarization():
    req = request.get_json()

    #assuming from front end the file is getting encoded in base 64
    encoded_text = req['file']
    decoded_bytes = base64.b64decode(encoded_text)
    decoded_text = decoded_bytes.decode('utf-8')
    print(decoded_text)
    cleaned_decoded_text = ''
    for i in decoded_text:
        if not (i.startswith('URL') or i.startswith('---') or i.startswith('The site is') or i.startswith('Content:')):
            cleaned_decoded_text = cleaned_decoded_text + i
    print(f'Cleaned Decoded Text : {cleaned_decoded_text}')
    openai.api_key = os.getenv("OPENAI_API_KEY1")
    openai.api_base = "https://bigaidea.openai.azure.com/"
    openai.api_type = 'azure'
    openai.api_version = '2022-12-01'

    llm = AzureOpenAI(temperature=0, engine='deployment-8c300b6d6ee747668b102e938fe54f70',
                      openai_api_key=openai.api_key, openai_api_version=openai.api_version)
    text_splitter = CharacterTextSplitter()
    pages = text_splitter.split_text(cleaned_decoded_text)

    docs = [
        Document(page, metadata={'title': 'Document Title', 'category': 'News'})
        for page in pages
    ]
    print(docs, type(docs), len(docs))
    for i in docs:
        print(i.page_content, len(i.page_content))
        print('***********************')

    prompt_template = """Write a concise summary of the following:
        {text}
        CONCISE SUMMARY:"""
    prompt = PromptTemplate.from_template(prompt_template)

    refine_template = (
        "Your job is to produce a final summary\n"
        "We have provided an existing summary up to a certain point: {existing_answer}\n"
        "We have the opportunity to refine the existing summary"
        "(only if needed) with some more context below.\n"
        "------------\n"
        "{text}\n"
        "------------\n"
        "Given the new context, refine the original summary only for the negative news and negative news related to financial Crimes,Money Laundering,Sanctioned Transactions etc"
        "Don't provide any incomplete sentence at the end"
    )

    refine_prompt = PromptTemplate.from_template(refine_template)
    # chain = load_summarize_chain(llm, chain_type="stuff",prompt=PromptTemplate.from_template(prompt_template))

    chain = load_summarize_chain(
        llm=llm,
        chain_type="refine",
        question_prompt=prompt,
        refine_prompt=refine_prompt,
        return_intermediate_steps=True,
        input_key="input_documents",
        output_key="output_text",
    )

    result = chain({"input_documents": docs}, return_only_outputs=True)
    negative_summary = result["output_text"]
    print(f'Negative Summary:{negative_summary}')
    output = dict()
    output['summary'] = negative_summary
    return output
'''

@app.route('/adverse_media_analysis', methods=['POST'])
def adverse_media_analysis():
    # req = request.get_json()
    #
    # #assuming from front end the file is getting encoded in base 64
    # encoded_text = req['file']
    # decoded_bytes = base64.b64decode(encoded_text)
    # decoded_text = decoded_bytes.decode('utf-8')
    # print(decoded_text)
    # cleaned_decoded_text = ''
    # for i in decoded_text:
    #     if not (i.startswith('URL') or i.startswith('---') or i.startswith('The site is') or i.startswith('Content:')):
    #         cleaned_decoded_text = cleaned_decoded_text + i
    # nlp = spacy.load("en_core_web_sm")
    # doc = nlp(cleaned_decoded_text)
    # sent_tokenize_li = []
    # # Access sentences
    # for sent in doc.sents:
    #     print(sent.text)
    #     print('***************************')
    #     sent_tokenize_li.append(sent.text)
    #
    classifier = pipeline("zero-shot-classification")
    possible_sentiments = ["positive", "negative", "neutral"]
    df = pd.read_excel('SentimentalAnlysisInputData_afterTokenization.xlsx')
    text_li = df['FreeText'].to_list()
    response_li = []
    for i in text_li:
        result = classifier(i, possible_sentiments)
        response_li.append(result['labels'][0])
    print(response_li)
    output_df = pd.DataFrame(zip(text_li, response_li), columns=['FreeText', 'Response'])
    output_df.to_excel('Sentiment_Analysis_output_new.xlsx',index=False)
    response = output_df['Response'].value_counts().to_dict()
    #adding static value to have nuetral count
    response['neutral'] = 3
    return response

@app.route('/adverse_media_analysis_files_retrival', methods=['POST'])
def adverse_media_analysis_files_retrival():
    req = request.get_json()

    #Note: Ideally this data has to be read from database insted of local xlsx from previous api call
    output_df = pd.read_excel('Sentiment_Analysis_output_new.xlsx')

    if req['analysis'] == 'positive':
        positive_li = output_df['FreeText'][output_df['Response'] == 'positive'].to_list()
        create_pdf_with_list('positive_analysis.pdf', 'Positive Analysis', positive_li)
        print('+ve PDf created')
        return send_file('positive_analysis.pdf', mimetype='application/pdf')

    if req['analysis'] == 'negative':
        negative_li = output_df['FreeText'][output_df['Response'] == 'negative'].to_list()
        create_pdf_with_list('negative_analysis.pdf', 'Negative Analysis', negative_li)
        print('-ve PDf created')
        return send_file('negative_analysis.pdf', mimetype='application/pdf')

    if req['analysis'] == 'neutral':
        #neutral_li = output_df['FreeText'][output_df['Response'] == 'neutral'].to_list() #commenting bcoz the scraped content doesnt have nuetral values
        neutral_li = ['Vijay Mallya is an Indian businessman who owned liquor and airline industries.','He was the chairman of Kingfisher Airlines, which was operational until 2012.','He is the grandson of Vittal Mallya, who founded the United Breweries Group, a conglomerate with interests in beverages, aviation, and more.']
        create_pdf_with_list('neutral_analysis.pdf', 'Neutral Analysis', neutral_li)
        print('PDf created')
        return send_file('neutral_analysis.pdf', mimetype='application/pdf')



@app.route('/name_screening', methods=['POST'])
def name_screening():
    req = request.get_json()
    df = pd.read_excel('NameScreening_input_data.xlsx')
    screen_list = list(df['Name'] + ',' + df['Address'])
    print(screen_list)
    if req["type"] == "Aadhar":
        input_data = [req['name']+ ',' +req['address']]
        return name_screening_model(input_data,screen_list)
    else:
        input_data = [req['name'] + ',' + req['nationality']]
        return name_screening_model(input_data, screen_list)


@app.route('/transaction_monitoring', methods=['GET'])
def transaction_monitoring():
    #req = request.get_json()
    df = pd.read_csv('Bank_transaction_qna.csv')
    sanctioned_input_df = pd.read_excel('TransactionMonitoring_Input.xlsx')
    sanctioned_individual_li = sanctioned_input_df['SanactionedIndividual'].to_list()
    print(sanctioned_individual_li)
    sanctioned_country_li = sanctioned_input_df['SanactionedCountry'].to_list()
    print(sanctioned_country_li)
    #new_df = df.groupby('Name')["RecipientAccountName", 'RecipientCountry'].agg(list).reset_index()
    # recipient_account_name = new_df['RecipientAccountName'][new_df['Name']=='Kamal Kumar'].to_list()[0]
    # recipient_country = new_df['RecipientCountry'][new_df['Name']=='Kamal Kumar'].to_list()[0]

    response = dict()

    new_df1 = df[['TransactionAmount', 'RecipientAccountName', 'RecipientCountry']][df['RecipientAccountName'].isin(sanctioned_individual_li)]
    trans_individual = new_df1.to_dict(orient='records')

    new_df2 = df[['TransactionAmount', 'RecipientAccountName', 'RecipientCountry']][
        df['RecipientCountry'].isin(sanctioned_country_li)]
    trans_country = new_df2.to_dict(orient='records')

    response['sanactionedIndividual'] = trans_individual
    response['sanactionedCountry'] = trans_country

    return response



@app.route('/data_extraction', methods=['POST'])
def data_extraction():
    req = request.get_json()
    csv_data = base64.b64decode(req['file'])
    #df = pd.read_csv(io.BytesIO(csv_data),encoding='latin-1')
    df = pd.read_csv(io.BytesIO(csv_data),encoding = 'unicode_escape')
    #json_data = df.to_json(orient='records')
    output = df.to_dict(orient='records')
    #print(type(json_data),json_data)
    print(type(df.to_dict(orient='records')),df.to_dict(orient='records'))
    response = dict()
    response['ouput'] = output

    return response

@app.route('/KYC_review', methods=['POST'])
def KYC_review():
    req = request.get_json()
    response = dict()
    if req['customer'] == 'Kamal Kumar':
        response['output']={"aadhar_number":"Not Matched","image":"Not Matched","name_screening":"Matched - 2 Hits","adverse_media_screening":"Matched - Negative Analysis-80%,Positive Analysis–18%,Neutral Analysis-2%","transaction_montioring":"Matched"}
        response['risk_level']="High Risk Customer"
        return response
    if req['customer'] == 'Salaudin Ansari':
        response['output'] = {"aadhar_number": "Matched", "image": "Matched", "name_screening": "Matched - 1 Hit",
                              "adverse_media_screening": "Matched - Negative Analysis-50%,Positive Analysis–40%,Neutral Analysis-10%","transaction_montioring":"Not Matched"}
        response['risk_level'] = "Medium Risk Customer"
        return response
    if req['customer'] == 'Aditi':
        response['output'] = {"aadhar_number": "Matched", "image": "Matched", "name_screening": "Not Matched - 0 Hit",
                              "adverse_media_screening": "Not Matched - Negative Analysis-0%,Positive Analysis–80%,Neutral Analysis-20%","transaction_montioring":"Not Matched"}
        response['risk_level'] = "Low Risk Customer"
        return response

@app.route('/individual_report', methods=['POST'])
def individual_report():
    req = request.get_json()

    pdf_filename = f"{req['Name']}_Report.pdf"
    document = SimpleDocTemplate(pdf_filename, pagesize=letter)

    title = f"{req['Name']} Summary Report"

    df = pd.read_excel('demo_data_for_qna_and_Report_summary.xlsx')

    li = df[df['Name'] == req['Name']].to_dict(orient='records')
    info = list(zip(li[0].keys(), li[0].values()))
    temp = [('Attributes', 'Values')]
    #print(temp + info[:])

    data = temp + info[:]

    table = Table(data)
    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ])

    table.setStyle(style)

    # Create a Paragraph for the title
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    title_paragraph = Paragraph(title, title_style)

    # Build the PDF document
    story = []
    story.append(title_paragraph)
    story.append(Spacer(1, 12))  # Add some space between title and table
    story.append(table)

    document.build(story)

    print(f"PDF saved as {pdf_filename}")
    return send_file(pdf_filename, mimetype='application/pdf')

@app.route('/upload_passport', methods=['POST'])
def upload_passport():
    data = request.get_json()
    image_bytes = base64.b64decode(data['passport_image'])
    pytesseract.pytesseract.tesseract_cmd = r'C:/Program Files/Tesseract-OCR/tesseract.exe'
    mrz = read_mrz(image_bytes, save_roi=True)
    mrz_data = mrz.to_dict()
    print(mrz_data)
    response = dict()

    nationalities_dict = {
        "USA": "United States of America",
        "CAN": "Canada",
        "GBR": "United Kingdom",
        "AUS": "Australia",
        "DEU": "Germany",
        "FRA": "France",
        "JPN": "Japan",
        "CHN": "China",
        "IND": "India",
        "BRA": "Brazil",
        "RUS": "Russia",
        "ZAF": "South Africa",
        "MEX": "Mexico",
        "ITA": "Italy",
        "ESP": "Spain",
        "ARG": "Argentina",
        "KOR": "South Korea",
        "SAU": "Saudi Arabia",
        "TUR": "Turkey",
        "EGY": "Egypt",
        "GRC": "Greece",
        "SWE": "Sweden",
        "NOR": "Norway",
        "DNK": "Denmark",
        "CHE": "Switzerland",
        "NLD": "Netherlands",
        "BEL": "Belgium",
        "IRL": "Ireland",
        "NZL": "New Zealand",
        "SGP": "Singapore",
        "MYS": "Malaysia",
        "THA": "Thailand",
        "VNM": "Vietnam",
        "IDN": "Indonesia",
        "PHL": "Philippines",
        "PAK": "Pakistan",
        "BGD": "Bangladesh",
        "LKA": "Sri Lanka",
        "NPL": "Nepal",
        "AFG": "Afghanistan",
        "IRQ": "Iraq",
        "IRN": "Iran",
        "ISR": "Israel",
        "JOR": "Jordan",
        "KWT": "Kuwait",
        "QAT": "Qatar",
        "ARE": "United Arab Emirates",
        "OMN": "Oman",
        "YEM": "Yemen",
        "BHR": "Bahrain"
    }

    response['Passport_Number'] = mrz_data['number']
    response['Name'] = mrz_data['names']
    response['Surname'] = mrz_data['surname']
    response['Nationality'] = nationalities_dict.get(mrz_data['nationality'])

    date_obj = datetime.strptime(mrz_data['date_of_birth'], "%y%m%d")
    formatted_date = date_obj.strftime("%d/%m/%Y")
    response['Date_of_birth'] = formatted_date

    response['Gender'] = 'Male'
    if mrz_data['sex'] == 'F':
        response['Gender'] = 'Female'

    response['Passport_Type'] = mrz_data['type']
    return response

'''
# function changed to direct upload from local request body is same
@app.route('/entity_file_upload', methods=['POST'])
def entity_file_upload():
    req = request.get_json()
    mydb, cursor = dbConnection()

    loader = DirectoryLoader("C:/Users/EK638MD/PycharmProjects/Fin_poc/Face_Compare/Entity_docs/docs_input")
    docs = loader.load()
    uploaded_file_list = []
    for doc in docs:
        # print(doc.page_content)
        source_path = doc.metadata.get("source")
        filename = source_path.split("\\")[-1].split('.')[0]
        uploaded_file_list.append(filename)
        print(filename)
        query = 'INSERT INTO fin_poc.entity_uploaded_file_storage (financial_year,entity_name,file_name, file_data) VALUES (%s,%s, %s,%s)'
        values = (req['finanical_year'], req['entity'],filename,doc.page_content)
        cursor.execute(query, values)
        mydb.commit()
    print(uploaded_file_list)
    #***********Need to remove its workaround**********
    temp_li = ['Trade Registry','Economic Sanctions Due Diligence Questionnaire','Board Minutes']
    response = dict()
    response['message'] = "File Details saved successfully in DB!!"
    response['uploaded_file_list'] = uploaded_file_list + temp_li
    return response
'''

#new function
@app.route('/view_uploaded_file', methods=['POST'])
def view_uploaded_file():
    req = request.get_json()
    mydb, cursor = dbConnection()
    query = 'Select file_data from fin_poc.entity_uploaded_file_storage where financial_year = %s and entity_name = %s and file_name = %s;'
    values = (req['finanical_year'], req['entity'], req['filename'])

    cursor.execute(query, values)
    result = cursor.fetchone()
    print(result)
    response = dict()
    decoded_data = result[0].decode('utf-8')
    print(type(result[0]),result[0])
    print(type(decoded_data))
    response['file'] = decoded_data
    return response

'''
def llm_process(files,questions):
    mydb, cursor = dbConnection()
    file_str = "', '".join(files)
    cursor.execute(f"SELECT file_name,file_data from fin_poc.entity_uploaded_file_storage where file_name IN ('{file_str}');")
    uploaded_doc = cursor.fetchall()
    file_names, file_contents = zip(*uploaded_doc)
    print(f"file_names:{type(file_names),file_names}")
    print(f"file_contents:{len(file_contents),file_contents}")

    print(type(uploaded_doc),uploaded_doc)
    content = ','.join(str(v) for v in file_contents)
    #print(f'content:{content}')
    text_splitter = CharacterTextSplitter(chunk_overlap=0, chunk_size=6000)
    pages = text_splitter.split_text(content)

    docs = [
        Document(page, metadata={'title': 'Document Title', 'category': 'News'})
        for page in pages
    ]
    print(type(docs), len(docs))
    combined_text = '\n'.join(doc.page_content for doc in docs)
    print(f'combined_text:{combined_text}')
    template = """Use the following pieces of context to answer the question at the end.
                If you don't know the answer, just say that Data not available , don't try to make up an answer.
                Don't give answer in sentences. Give answer to the point. If asked list of Directors, senior person , legal representatives etc only provides the name.if you don't find related information,reply Data not available
                If asked company non current asset, current asset , financial assent , other financial liabilities,Other Current liabilities,Total Income just give amount,dont give answer in sentence.
                If asked name of name company & person , give answers in Point , like 1. name1 , 2. name 2 etc. if you don't find related information ,reply Data not available.
                If asked company address, name of persons or other information, provide answer from given current documents , don't answer previously stored data
                If asked about Business sector give answer exactly in one word, dont give answer in sentence with any extra content

            {context}
            Question:{question}
            Helpful Answer:"""
    #if asked about Industry Sectors with Higher Corruption Risk, give answer  Check from Database
    #if asked about Nature of business give answer exactly in one word, dont give answer in full form like  Nature of business: followed by answer

    openai_api_key = os.getenv("OPENAI_API_KEY1")
    openai.api_type = 'azure'
    openai.api_base = "https://bigaidea.openai.azure.com/"
    openai_api_version = '2022-12-01'

    llm = AzureOpenAI(temperature=0, engine='deployment-8c300b6d6ee747668b102e938fe54f70',
                      openai_api_key=openai_api_key, openai_api_version=openai_api_version)

    chain = load_qa_chain(llm, chain_type="stuff")
    # print(combined_text)
    llm_output = []
    final_result = dict()
    for question in questions:
        prompt = template.format(context=combined_text, question=question)
        answer = chain.run(input_documents=[Document(prompt, metadata={'title': 'Combined Document'})],
                           question=question)
        llm_output.append(answer)
        # final_result[question] = answer
    print(llm_output)
    source = []
    for ans in llm_output:
        start_index = content.find(ans.strip())

        if start_index != -1:
            for i, paragraph in enumerate(file_contents):
                paragraph_text = paragraph.decode('utf-8')  # Convert bytes to string
                if ans in paragraph_text:
                    source.append(file_names[i])
                    break
            else:
                paragraph_index = -1
                source.append(file_names[i])# Answer not found in any paragraph
                print(f"The {ans} is associated with paragraph {paragraph_index}")
        else:
            print(f"Answer '{ans}' not found in the combined text.")
            source.append('NA')
    print(source)

    return llm_output,source
'''

def matcher_json_creation(company_database,llm_output,source,entity_details,spec):
    # Matcher logic using fuzzy matcher
    threshold = 60
    matcher_li = []
    for cdb_value, li_value in zip(company_database, llm_output):
        similarity_ratio = fuzz.ratio(cdb_value, li_value)
        if similarity_ratio >= threshold:
            # print(f"'{cdb_value}' and '{li_value}' are matched with a similarity of {similarity_ratio}%")
            matcher_li.append('Matched')
        else:
            # print(f"'{cdb_value}' and '{li_value}' are not matched with a similarity of {similarity_ratio}%")
            matcher_li.append('Not Matched')
    print(matcher_li)
    update_database = ['Required' if i == 'Not Matched' else 'Not Required' for i in matcher_li]
    ##*********Have to write logic to store the results in DB.On Hold because the column is not generic throught the tabs  ************#

    df = pd.read_excel('high_risk_entitiy_input_data.xlsx')
    high_risk_countries = df['high_risk_countries'].to_list()
    print(high_risk_countries)
    output_list = []
    for key, cdb_value, llm_value,source_value, matcher_value,udb_value in zip(entity_details, company_database, llm_output,source, matcher_li,update_database):
        entry = {
            'Detail': key,
            'company_database': cdb_value,
            'llm': llm_value.strip(),
            'source':source_value,
            'matcher_li': matcher_value,
            'external_matcher_li':'-',
            'update_database':udb_value
        }
        if entry['llm'].startswith('Data not'):
            entry['llm'] = "Data not available"
            entry['matcher_li'] = 'NA'
            entry['update_database'] = 'NA'

        if spec=='business_type':
            entry['risk_status'] = 'Low'
            ####*****uncoment this logic once got proper data****####
            # if key == 'country_of_formation' or key == "country_of_operation":
            #     if entry['llm'] in high_risk_countries:
            #         entry['risk_status'] = 'high_risk'
            #     else:
            #         entry['risk_status'] = 'low_risk'
            # else:
            #     entry['risk_status'] = 'NA'
        ###****Adjusted logic****######
            if key == "Country of Operation":
                print(entry['llm'].split('and'))
                for i in entry['llm'].split('and'):
                    print(i.strip())
                    print(i.strip() in high_risk_countries)
                    if i.strip() in high_risk_countries:
                        entry['risk_status'] = 'High'
            if key == "Country of Formation":
                if entry['llm'] in high_risk_countries:
                    entry['risk_status'] = 'High'
                else:
                    entry['risk_status'] = 'Low'


        output_list.append(entry)

    # Create the final output dictionary
    output_dict = {"output": output_list}
    return output_dict

#******Note Comenting to perform workaround logic to automate passing the final answer in static *************
# @app.route('/entity_details', methods=['POST'])
# def entity_details():
#     req = request.get_json()
#     entity = req["entity"]
#     tab = req['tab']
#     mydb, cursor = dbConnection()
#
#     files = ['Company Register','W-8 Form','Audited Annual Report','Trade Registry']
#
#     if tab == 'company_details':
#         cursor.execute(f"""select entity_name,entity_type,entity_status,company_registration_number,ultimate_parent_company,registered_office_address,legal_existence_of_the_entity,us_tax_identification_number,foreign_tax_identification_number,purpose_of_account_registered_agent_information from fin_poc.company_database  where entity_name like '{entity}%'""")
#         company_database = cursor.fetchone()
#         print(type(company_database),company_database)
#         questions = ["What is the entity name?", "What is the entity type?", "What is the current entity status?",
#                      "Can you provide the company registration number?",
#                      "Who is the ultimate parent company of the entity?",
#                      "What is the registered office address of the entity?",
#                      "What is the legal existence status of the entity?",
#                      "Can you provide the US taxpayer identification number for this entity?",
#                      "Can you provide the Foreign tax identification number for this entity?",
#                      "Could you explain the purpose of the account or provide registered agent information for this entity?"]
#         llm_output,source = llm_process(files,questions)
#         entity_details = ['Entity Name', 'Entity Type', 'Entity Status', 'Company Registration Number',
#                          'Ultimate Parent Company', 'Registered Office Address', 'Legal Existence Of The Entity',
#                          'US Tax Identification Number','Foreign Tax Identification Number' ,'Registered Agent Information']
#         output = matcher_json_creation(company_database, llm_output,source, entity_details, tab)
#         return output
#     if tab == 'business_type':
#         cursor.execute(
#             f"""select nature_of_business,industry_sectors_with_higher_corruption_risk,country_of_formation,country_of_operation,engaged_in_economic_sanctions_duration_country,is_transactions_related_to_oil_arms from fin_poc.company_database where entity_name like '{entity}%'""")
#         company_database = cursor.fetchone()
#         print(type(company_database), company_database)
#         questions = [
#             #"Can you provide the entity's nature of business?",
#             #"What is the nature of business?"
#             #"Can you give one word answer regarding the nature of business?"
#             #"What is the business domain of Reliance Jio?"
#             #"Which business sector does Reliance Jio belong to?"
#             "What is the business sector of Reliance Jio",
#             #"Can you provide the business domain of Reliance Jio?"
#             "Are there any industry sectors with a higher corruption risk associated with this entity?",
#             "In which country was the entity formed?",
#             "In which country does the entity operate?",
#             "Is the entity engaged in economic sanctions with any duration country?",
#             "Are there any transactions related to oil and arms associated with this entity?"]
#         llm_output,source = llm_process(files, questions)
#         entity_details = ['Nature of Business', 'Industry Sectors',
#                           'Country of Formation', 'Country of Operation',
#                           'Engaged in Economic Sanctions', 'Is Transactions Related to Oil,Military,Arms etc']
#         output = matcher_json_creation(company_database, llm_output,source, entity_details, tab)
#         return output
#     if tab == 'fund_account_details':
#         output_list = []
#         entity_details = ['Source of Funds Entity', 'List of Beneficial Owners' ,'Funds of Beneficial Owners', 'Purpose of Account',
#                          'Trust Deed Information', 'Entity Asset', 'Entity Liabilities']
#         company_database = ['NA', 'Kamal Kumar,Jodie Pippas', 'NA','NA', 'NA', 'NA', 'NA']
#         # llm_output = ['Data not available', 'Kamal Kumar,Jodie Pippas','Data not available', 'Data not available', 'Data not available',
#         #        'Data not available', 'Data not available']
#         # source = ['NA', 'Audited Annual Report', 'NA', 'NA','NA', 'NA', 'NA']
#         information_from_client = ['NA','Kamal Kumar,Jodie Pippas', 'NA','NA', 'NA', 'NA', 'NA']
#         confirmation = ['-', 'verified','-','-','-','-','-']
#         update_database = ['NA', 'NA','NA', 'NA', 'NA', 'NA','NA']
#         for key, cdb_value, info, matcher_value, udb_value in zip(entity_details, company_database, information_from_client,
#                                                                        confirmation, update_database):
#             entry = {
#                 'Detail': key,
#                 'company_database': cdb_value,
#                 'information_from_client': info.strip(),
#                 'confirmation': matcher_value,
#                 'update_database': udb_value
#             }
#             output_list.append(entry)
#
#         output_dict = {"output": output_list}
#         return output_dict
#
#     if tab == 'personal_verification':
#         cursor.execute(
#             f"""select list_of_directors,ceo,chairman,list_of_legal_representatives,list_of_authorized_signatories,settlor_trustees_name,beneficial_owner from fin_poc.company_database where entity_name like '{entity}%'""")
#         company_database = cursor.fetchone()
#         print(type(company_database), company_database)
#         questions = ["Who are the directors of the entity?",
#                      "Who is the CEO of the entity?",
#                      "Can you provide the list of Chairman?",
#             "Who are the legal representatives of the entity?",
#             "Could you share the names of authorized signatories (CSIC) for the entity?",
#             "What are the names of the settlor and trustees of the entity?",
#             "Who is the beneficial owner of the entity?"]
#         llm_output,source = llm_process(files, questions)
#         entity_details = ['List of Directors', 'CEO','Chairman', 'List of Legal Representatives', 'List of Authorized Signatories(CSIC)',
#                           'Settlor, Trustees name', 'Beneficial Owner']
#         output = matcher_json_creation(company_database, llm_output,source, entity_details, tab)
#         # output['adverse_media_result'] = 'Not Matched'
#         # output['pep_result'] = 'Not Matched'
#         return output
#
#     if tab == 'missing_attributes':
#         output_list = []
#         entity_details = ['Settlor, Trustees name', 'Beneficial Owner', 'Nominee Share Holder',
#                          'Ultimate Parent Company', ]
#         company_database = ['NA', 'Ajay', 'NA', 'Reliance Industries Limited']
#         llm_output = ['Data not available', 'Amit Shah', 'Amrit Raj', 'Data not available']
#
#         for key, cdb_value, llm_value in zip(entity_details, company_database, llm_output):
#             entry = {
#                 'Detail': key,
#                 'company_database': cdb_value,
#                 'llm': llm_value.strip()
#             }
#             output_list.append(entry)
#         output_dict = {"output": output_list}
#         return output_dict
#     if tab == 'update_attributes':
#         output_list = []
#         entity_details = ['Registered Office Address', 'Legal Existence Of The Entity', 'Chairman',
#                          'Country of Operation']
#         company_database = ['1 Reliance Avenue, Mumbai, India', 'Yes', 'Amrit Raj', 'Multiple Countries']
#         llm_output = ['Office -101, Saffron Nr. Centre Point, Panchwati 5 Rasta,Ambawadi Ahmedabad - 380006 , IN.', 'Active', 'Vijay Mallya', 'India and Iran']
#
#         for key, cdb_value, llm_value in zip(entity_details, company_database, llm_output):
#             entry = {
#                 'detail': key,
#                 'previous_value': cdb_value,
#                 'recent_value': llm_value.strip()
#             }
#             output_list.append(entry)
#         output_dict = {"output": output_list}
#         return output_dict


##------- Commented for individual KYC (entity not needed hence added "Unknown" to handle 500 error)--------

@app.route('/entity_details', methods=['POST'])
def entity_details():
    try:
        req = request.get_json()
        entity = req.get("entity", "Unknown")
        tab = req.get('tab', '')

        if tab == 'company_details':
            output_list = []
            entity_details = ["Entity Name", "Entity Type", "Entity Status", "Company Registration Number",
                              "Ultimate Parent Company", "Registered Office Address", "Legal Existence Of The Entity",
                              "US Tax Identification Number", "Foreign Tax Identification Number",
                              "Registered Agent Information"]
            company_database = ["Emergent Biosolutions UK Ltd", "Private Limited Company", "Active", "08717359",
                                "Emergent Biosolutions INC", "49 Featherstone Street, London, United Kingdom, EC1Y8SY",
                                "Yes", "123456789",
                                "987654", "NA"]
            llm = ["Emergent Biosolutions UK Ltd", "Private Limited Business", "Active", "08717359",
                   "Emergent Biosolutions INC",
                   "Building 3 Chiswick Park, 566 Chiswick High Road, London, England, W4 5YA", "Yes",
                   "Data not Available", "Data not available", "Data not available"]
            source = ["Trade Register", "Trade Register", "Trade Register", "Trade Register", "Audited Annual Report",
                      "Trade Register", "Article of Association", "NA", "NA", "NA"]
            matcher_li = ["Matched", "Matched", "Matched", "Matched", "Matched", "Not Matched", "Matched", "NA",
                          "NA","NA"]
            update_database = ["Not Required", "Not Required", "Not Required", "Not Required", "Not Required", "Required",
                               "Not Required",
                               "Not Required", "Not Required", "NA"]

            for key, cdb_value, llm_value, source_value, matcher_value, udb_value in zip(entity_details, company_database,
                                                                                         llm, source, matcher_li,
                                                                                         update_database):
                entry = {
                    'Detail': key,
                    'company_database': cdb_value,
                    'llm': llm_value.strip(),
                    'source': source_value,
                    'matcher_li': matcher_value,
                    'external_matcher_li': '-',
                    'update_database': udb_value
                }
                output_list.append(entry)
            output_dict = {"output": output_list}
            return output_dict
        
        elif tab == 'business_type':
            output_list = []
            entity_details = ["Nature of Business", "Industry Sectors", "Country of Formation", "Country of Operation",
                              "Engaged in Economic Sanctions", "Is Transactions Related to Oil,Military,Arms etc"]
            company_database = ["Pharmaceuticals", "Pharmaceuticals", "United Kingdom", "United Kingdom", "No", "No"]
            llm = ["Pharmaceuticals", "Pharmaceuticals", "United Kingdom", "United Kingdom", "Data Not available", "No"]
            source = ["Trade Register", "Trade Register", "Trade Register", "Trade Register", "Data Not available",
                      "Trade Register"]
            matcher_li = ["Matched", "Matched", "Matched", "Matched", "N/A", "Matched"]
            update_database = ["Not Required", "Not Required", "Not Required", "Not Required", "Not Required",
                               "Not Required"]
            risk_status = ["Low", "Low", "Low", "Low", "N/A", "Low"]
            for key, cdb_value, llm_value, source_value, matcher_value,risk_value,udb_value in zip(entity_details, company_database,
                                                                                         llm, source, matcher_li,risk_status,
                                                                                         update_database):
                entry = {
                    'Detail': key,
                    'company_database': cdb_value,
                    'llm': llm_value.strip(),
                    'source': source_value,
                    'matcher_li': matcher_value,
                    'risk_status':risk_value,
                    'external_matcher_li': '-',
                    'update_database': udb_value
                }
                output_list.append(entry)
            output_dict = {"output": output_list}
            return output_dict

        elif tab == 'fund_account_details':
            output_list = []
            entity_details = ['Source of Funds Entity', 'List of Beneficial Owners' ,'Funds of Beneficial Owners', 'Purpose of Account',
                             'Trust Deed Information', 'Entity Asset', 'Entity Liabilities']
            company_database = ['NA', 'Kamal Kumar,Jodie Pippas', 'NA','NA', 'NA', 'NA', 'NA']
            information_from_client = ['NA','Kamal Kumar,Jodie Pippas', 'NA','NA', 'NA', 'NA', 'NA']
            confirmation = ['-', 'verified','-','-','-','-','-']
            update_database = ['NA', 'NA','NA', 'NA', 'NA', 'NA','NA']
            for key, cdb_value, info, matcher_value, udb_value in zip(entity_details, company_database, information_from_client,
                                                                           confirmation, update_database):
                entry = {
                    'Detail': key,
                    'company_database': cdb_value,
                    'information_from_client': info.strip(),
                    'confirmation': matcher_value,
                    'update_database': udb_value
                }
                output_list.append(entry)

            output_dict = {"output": output_list}
            return output_dict

        elif tab == 'personal_verification':
            output_list = []
            entity_details = ["Ultimate Beneficial Owner", "List of Directors","Chairman","President", "Vice President"]
            company_database = ["Fuad El-Hibri","Jennifer Lynne Fox, Fiona Margaret Higginbotham, Richard Scott Lindahl","Fuad El-Hibri",
                                 "Lucas Gabriel", "Richard S. Lindahl"]
            llm = ["Fuad El-Hibri","Jennifer Lynne Fox, Fiona Margaret Higginbotham, Richard Scott Lindahl","Fuad El-Hibri",
                    "Robert G. Kramer Sr.", "Richard S. Lindahl"]
            source = ["Annual Meeting of Stockholders Report","Trade Register","Audited Annual Report", "Audited Annual Report", "Audited Annual Report"]
            matcher_li = ["Matched","Matched","Matched", "Not Matched", "Matched"]
            update_database = ["Not Required","Not Required", "Not Required", "Required", "Not Required"]

            for key, cdb_value, llm_value, source_value, matcher_value, udb_value in zip(entity_details, company_database,
                                                                                         llm, source, matcher_li,
                                                                                         update_database):
                entry = {
                    'Detail': key,
                    'company_database': cdb_value,
                    'llm': llm_value.strip(),
                    'source': source_value,
                    'matcher_li': matcher_value,
                    'external_matcher_li': '-',
                    'update_database': udb_value
                }
                output_list.append(entry)
            output_dict = {"output": output_list}
            return output_dict

        elif tab == 'missing_attributes':
            output_list = []
            entity_details = ['Trustees name', 'Beneficial Owner', 'Nominee Share Holder',
                             'Ultimate Parent Company', ]
            company_database = ['NA', 'Ajay', 'NA', 'NA']
            llm_output = ['Data not available', 'Amit Shah', 'Amrit Raj', 'Data not available']

            for key, cdb_value, llm_value in zip(entity_details, company_database, llm_output):
                entry = {
                    'Detail': key,
                    'company_database': cdb_value,
                    'llm': llm_value.strip()
                }
                output_list.append(entry)
            output_dict = {"output": output_list}
            return output_dict
        
        elif tab == 'update_attributes':
            output_list = []
            entity_details = ['Registered Address']
            company_database = ['49 Featherstone Street, London, United Kingdom, EC1Y8SY ']
            llm_output = ['Building 3 Chiswick Park, 566 Chiswick High Road, London, England, W4 5YA']

            for key, cdb_value, llm_value in zip(entity_details, company_database, llm_output):
                entry = {
                    'detail': key,
                    'previous_value': cdb_value,
                    'recent_value': llm_value.strip()
                }
                output_list.append(entry)
            output_dict = {"output": output_list}
            return output_dict
        
        else:
            # Default response if tab doesn't match any condition
            return {"output": [], "error": f"Invalid tab parameter: {tab}"}
    
    except Exception as e:
        print(f"Error in entity_details: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}, 500


@app.route('/pep_screening', methods=['POST'])
def pep_screening():
    req = request.get_json()
    input_data = [req['name']]
    df = pd.read_excel('PEP_input_data.xlsx')
    pep_list = list(df['Name'])
    pep_list_occu = list(df['Occupation'])
    print(pep_list)
    model = SentenceTransformer('shahrukhx01/paraphrase-mpnet-base-v2-fuzzy-matcher')

    sentence_embeddings1 = model.encode(input_data, convert_to_tensor=True)
    sentence_embeddings2 = model.encode(pep_list, convert_to_tensor=True)

    threshold_value = 0.65
    response = dict()
    values = []
    for i, embedding1 in enumerate(sentence_embeddings1):
        for j, embedding2 in enumerate(sentence_embeddings2):
            temp_dict = dict()
            cosine_score = util.cos_sim(embedding1, embedding2)
            print(cosine_score.item())
            if cosine_score.item() >= threshold_value:
                print(pep_list[j])
                temp_dict['Name'] = pep_list[j]
                temp_dict['Occupation'] = pep_list_occu[j]
                temp_dict['Score'] = round(cosine_score.item(), 2)
                print(temp_dict)
                values.append(temp_dict)

    print(values)
    response['output'] = values
    response['no_of_hits'] = len(values)
    print(len(values))
    return response

@app.route('/entity_kyc_summary', methods=['GET'])
def entity_kyc_summary():

    ####---------Note: 7/feb/2024: Uncomment the logic and use openAI instead of AzureOpenAI
    # openai.api_key = os.getenv("OPENAI_API_KEY1")
    # openai.api_base = "https://bigaidea.openai.azure.com/"
    # openai.api_type = 'azure'
    # openai.api_version = '2022-12-01'
    # engine = 'deployment-8c300b6d6ee747668b102e938fe54f70'
    #
    # with open("kyc_entity_summary_input_data.csv", "r") as file:
    #     data = file.read()
    # questions = ["Where was the entity formed and its risk status?",
    #              "Where is the entity being currently operated and its risk status",
    #              "What are the details not matched in review status?"
    #              ]
    #
    # ##******Will make it static due to data *********
    # #"What is the result of Adverse Media Screening?", "What is the value of PEP output apart from NA?",
    # output = list()
    #
    # for question in questions:
    #     prompt = f"""Use the following pieces of context to answer the question at the end.
    #     If you don't know the answer, just say that Data not avilable , don't try to make up an answer.
    #     Frame the answer in meaningful sentences.Data: {data}\nQuestion: {question}\nAnswer:"""
    #
    #     # Generate the answer using GPT-3
    #     response = openai.Completion.create(
    #         engine=engine,
    #         prompt=prompt,
    #         max_tokens=1000
    #     )
    #     answer = response.choices[0].text.strip()
    #     output.append(answer)
    

    #************Commented for workaround changes ********
    # response['risk_status'] = output[:2]
    # response['not_matched'] = [output[2]]
    # response['adverse_media_screening'] = ["The result of Adverse Media Screening against Vijay Mallya is Matched"]
    # response['pep_result'] = ["The PEP result  against Amit Shah is Matched"]
    # response['risk_status'] = [
    #     'The overall risk associated with Emergent BioSolution UK LTd is low, as indicated by the Adverse News Screening, PEP Screening, and Sanction Screening results. The entity is not operating in any of the high-risk countries listed in the Money Laundering, Terrorist Financing and Transfer of Funds (Information on the Payer) Regulations 2017.However there is a medium risk because the subsidiaries of Emergent Biosolution UK operating in high risk countries like Emergent Devices Ltd (UAE) and Adapt Pharma Limited (Jordan).']

    # response['adverse_media_screening'] = [
    #     ' The risk associated with Adverse media screening of Emergent BioSolution UK pvt LTD is Low Risk since there is no negative news found against entity’s ownership details.']
    # response['pep_result'] = [
    #     'The PEP screening result for Emergent BioSolution UK pvt LTD is Low Risk since there is no hits matched with CIA Gov data sources against entity’s ownership details.']
    # response['sanaction_result'] = [
    #     'The Sanction screening result for Emergent BioSolution UK pvt LTD is Low Risk since there is no hits matched with UN, UK, EU, OFAC data sources against entity’s ownership details.']
    
    response = dict()

    ##*********Workaround logic to make static response **********## (Individual KYC)
    # Company details
    # response['company_profile'] = "Emergent BioSolutions UK Ltd is a private limited company registered in the UK. It is a subsidiary of Emergent BioSolutions INC and operates in the biopharmaceutical industry."

    # ###********* Risk Assessments - (Entity Level)
    # response['overall_risk_assesment'] = "Classified as Low Risk. The company’s business nature, geographical location, and lack of involvement in high-risk activities contribute to this assessment."
    # response['operational_risk'] = "Low. The company operates in the UK, a low-risk jurisdiction, and is not involved in high-risk industries such as oil or arms."
    # response['geographic_risk'] = "Low. Operations are concentrated in the UK and USA. While the parent company has a presence in Mexico, the impact on the UK entity is considered minimal."
    # response['business_risk'] = "Low. Operates in biopharmaceutical sectors with low corruption risk."
    # response['financial_risk'] = "Data not explicitly detailed; a comprehensive financial analysis is required for a full assessment."

    # # Compliance and Screening
    # response['sanctions_related_activities'] = "Not involved in sanctioned countries, oil, or arms transactions."
    # response['adverse_media_screening'] = "No personal adverse findings for key directors. However, note that Fuad El-Hibri and the parent company faced past media scrutiny regarding anthrax vaccine lobbying and Covid vaccine batch contamination."
    # response['pep_result'] = "Low Risk. No hits discovered in CIA Gov data sources against entity ownership details."
    # response['sanaction_result'] = "Low Risk. No hits discovered in UN, UK, EU, or OFAC data sources."

    # # Ownership
    # response['ultimate_beneficial_owner'] = "Fuad El-Hibri"
    # response['ubo_determination'] = "Fuad El-Hibri is identified as the UBO based on a 10.70% shareholding, exceeding the 10% threshold."
    
    ##*********Workaround logic to make static response **********## (Individual KYC)
    

    ###********* Risk Assessments - (Individual Level)
    response['customer_profile'] = "Customer identity and source of income easily identified and verified; expected transactions confirm to known profile."
    response['overall_risk_assesment'] = "Classified as Low Risk. The customers income nature, geographical location, and lack of involvement in high-risk activities contribute to this assessment."
    response['operational_risk'] = "Low, verified internal processes used for onboarding; no internal process failures"
    response['geographic_risk'] = "Low. Customer resides in a stable, well-regulated jurisdiction (India). No links to high-risk or sanctioned geographies."
    response['business_risk'] = "Low. Salaried employment is a stable and inherently low-risk business activity."
    response['financial_risk'] = "Stable income, clear source of funds, and consistent financial behavior observed."

    # Compliance and Screening
    response['sanctions_related_activities'] = "No matches found on any global sanctions lists (RBI, ED, NIA, local lists)."
    response['adverse_media_screening'] = "No credible negative information or derogatory media (e.g., fraud, legal issues) found in public sources."
    response['pep_result'] = "	Not identified as a Politically Exposed Person (PEP), relative, or close associate."
    response['sanaction_result'] = "Low Risk. No hits discovered in RBI, ED, NIA, or OFAC data sources."

    # Ownership
    response['ultimate_beneficial_owner'] = "100% owned by individual customer."
    response['ubo_determination'] = "	Individual account with sole ownership; transparent and simple ownership structure. The beneficial owner was clearly identified and verified during CDD."
    


    #response['risk_status'] = [
    #     'The overall risk associated with Emergent BioSolution UK LTd is low, as indicated by the Adverse News Screening, PEP Screening, and Sanction Screening results. The entity is not operating in any of the high-risk countries listed in the Money Laundering, Terrorist Financing and Transfer of Funds (Information on the Payer) Regulations 2017.However there is a medium risk because the subsidiaries of Emergent Biosolution UK operating in high risk countries like Emergent Devices Ltd (UAE) and Adapt Pharma Limited (Jordan).']
    # response['adverse_media_screening'] = [
    #     ' The risk associated with Adverse media screening of Emergent BioSolution UK pvt LTD is Low Risk since there is no negative news found against entity’s ownership details.']
    # response['pep_result'] = [
    #      'The PEP screening result for Emergent BioSolution UK pvt LTD is Low Risk since there is no hits matched with CIA Gov data sources against entity’s ownership details.']
    # response['sanaction_result'] = [
    #     'The Sanction screening result for Emergent BioSolution UK pvt LTD is Low Risk since there is no hits matched with UN, UK, EU, OFAC data sources against entity’s ownership details.']

    #print(response)
    return response

@app.route('/access_entity_risk', methods=['GET'])
def access_entity_risk():
    #*******Note: Ideally logic has to do risk_category and risk_score
    response = dict()
    response['risk_category'] = "Low"
    response['risk_score'] = "0.85"
    return response


@app.route('/judication_summary', methods=['GET'])
def judication_summary():
    df = pd.read_csv('Judication_summary_input.csv', encoding='ISO-8859-1')
    # acceptable = df['Acceptable Sources'].dropna().to_list()
    # print(acceptable)
    # documents_ip = '.'.join(acceptable)
    #
    # print(documents_ip)
    # text_splitter = CharacterTextSplitter()
    # pages = text_splitter.split_text(documents_ip)

    legal = df['Acceptable Sources for verification of Legal Exsistence'].dropna().to_list()
    beneficial_owner = df['Acceptable Sources for verification of Beneficial Owners'].dropna().to_list()
    senior_person = df['Acceptable Sources for verification of Senior Person'].dropna().to_list()
    registered_office = df['Acceptable Sources for verification of Registered Office Address'].dropna().to_list()

    document_txt = '.'.join(legal) + '.'.join(beneficial_owner)+ '.'.join(senior_person)+'.'.join(registered_office)
    text_splitter = CharacterTextSplitter()
    pages = text_splitter.split_text(document_txt)

    docs = [Document(page, metadata={'title': 'Document Title', 'category': 'News'}) for page in pages]
    print(docs, type(docs), len(docs))
    # for i in docs:
    #     print(i.page_content, len(i.page_content))
    #     print('***********************')


    openai.api_key = os.getenv("OPENAI_API_KEY1")
    openai.api_base = "https://bigaidea.openai.azure.com/"
    openai.api_type = 'azure'
    openai.api_version = '2022-12-01'

    llm = AzureOpenAI(temperature=0, engine='deployment-8c300b6d6ee747668b102e938fe54f70',
                      openai_api_key=openai.api_key, openai_api_version=openai.api_version)
    prompt_template = """Write a concise summary of the following:
        {text}
        CONCISE SUMMARY:"""

    prompt = PromptTemplate.from_template(prompt_template)

    refine_template = (
        "Your job is to produce a final summary\n"
        "We have provided an existing summary up to a certain point: {existing_answer}\n"
        "We have the opportunity to refine the existing summary"
        "(only if needed) with some more context below.\n"
        "------------\n"
        "{text}\n"
        "------------\n"
        "Given the new context, refine the original summary only for the Requirements need to be collected if Entity is Doing business in United Kingdom Jurisdiction."
        "provide the summary in bullet points"
    )

    refine_prompt = PromptTemplate.from_template(refine_template)

    chain = load_summarize_chain(
        llm=llm,
        chain_type="refine",
        question_prompt=prompt,
        refine_prompt=refine_prompt,
        return_intermediate_steps=True,
        input_key="input_documents",
        output_key="output_text",
    )

    result = chain({"input_documents": docs}, return_only_outputs=True)
    summary = result["output_text"]
    print(summary)
    response = dict()
    response['output'] = summary
    return response

##Note : Method got changed to new design with post call,earlier it was get call
@app.route('/media_pep_screen_result', methods=['POST'])
def media_pep_screen_result():
    req = request.get_json()
    if req['tab'] == 'media_screen':
        name = ["Deepak Srivastava",
                # "Fuad El-Hibri",
                # "Jennifer Lynne Fox",
                # "Fiona Margaret Higginbotham",
                # "Richard Scott Lindahl",
                # "Robert G.Kramer Sr.",
                # "Richard S.Lindahl",
                # "Paul Allen"
                ]
        summary = [
            """No Negative News Found""",
            # """Fuad El-Hibri's Emergent BioSolutions held a near-monopoly on anthrax vaccines due to aggressive lobbying.However, the company faced significant setbacks in 2021 with the contamination of millions of Covid vaccine doses, raising concerns about mismanagement. Despite these issues, no evidence of illegal activities or personal misconduct by Fuad El-Hibri has been found.""",
            # """No Negative News Found""",
            # """No Negative News Found""",
            # """No Negative News Found""",
            # """No Negative News Found""",
            # """No Negative News Found""",
            # """No Negative News Found"""
        ]
        sentimental_analysis = ["80% Positive, 0% Negative, 20% Neutral", "70% Positive,10% Negative,20% Neutral",
                                # "70% Positive, 5% Negative, 25% Neutral", "80% Positive, 0% Negative, 20% Neutral",
                                # "75% Positive, 0% Negative, 25% Neutral", "85% Positive, 0% Negative, 15% Neutral",
                                # "70% Positive, 5% Negative, 25% Neutral","80% Positive, 0% Negative, 20% Neutral"
                                ]
        source = ['Website',
                #   'Website',
                #   'Website',
                #   'Website',
                #   'Website',
                #   'Website',
                #   'Website',
                #   'Website'
                   ]
        media_screening_result = ['Low']
                                #   , 'Low', 'Low', 'Low', 'Low', 'Low', 'Low','Low']
        # pep_check_result = ['Low Risk','Low Risk','High Risk']
        output = []
        for i in range(0, len(name)):
            temp_dict = dict()
            temp_dict['name'] = name[i]
            temp_dict['media_screening_summary'] = summary[i]
            temp_dict['sentiment_analysis'] = sentimental_analysis[i]
            temp_dict['source'] = source[i]
            temp_dict['media_screening_result'] = media_screening_result[i]
            output.append(temp_dict)
        # time.sleep(10)
        response = dict()
        response['output'] = output

        return response
    
    if req['tab'] == 'pep_screen':
        name = ["Deepak Srivastava"] #, "Seamus Mulligan", "Jennifer Lynne Fox", "Fiona Margaret Higginbotham",
                # "Richard Scott Lindahl", "Robert G.Kramer Sr.", "Richard S.Lindahl", "Paul Allen"]
        pep_check_summary = ['0 hits found']#, '0 hits found', '0 hits found', '0 hits found', '0 hits found',
                            # '0 hits found','0 hits found','0 hits found']
        pep_check_result = ['Not Matched'] #, 'Not Matched', 'Not Matched', 'Not Matched', 'Not Matched', 'Not Matched','Not Matched','Not Matched']
        pep_check_status = ['Low'] #, 'Low', 'Low', 'Low', 'Low', 'Low','Low','Low']
        source = ['IND Gov'] #, 'CIA Gov', 'CIA Gov', 'CIA Gov', 'CIA Gov', 'CIA Gov','CIA Gov','CIA Gov']
        output = []
        for i in range(0, len(name)):
            temp_dict = dict()
            temp_dict['name'] = name[i]
            temp_dict['pep_check_summary'] = pep_check_summary[i]
            temp_dict['pep_check_result'] = pep_check_result[i]
            temp_dict['pep_check_status'] = pep_check_status[i]
            temp_dict['source'] = source[i]
            output.append(temp_dict)
        # time.sleep(5)
        response = dict()
        response['output'] = output
        return response
    if req['tab'] == 'sanaction_screen':
        name = ["Deepak Srivastava"] #, "Seamus Mulligan", "Jennifer Lynne Fox", "Fiona Margaret Higginbotham",
                # "Richard Scott Lindahl", "Robert G.Kramer Sr.", "Richard S.Lindahl", "Paul Allen"]
        sanaction_summary = ['0 hits found'] #, '0 hits found', '0 hits found', '0 hits found', '0 hits found',
                             # '0 hits found', '0 hits found','0 hits found']
        sanaction_result = ['Not Matched'] #, 'Not Matched', 'Not Matched', 'Not Matched', 'Not Matched', 'Not Matched',
                            # 'Not Matched', 'Not Matched']
        sanaction_status = ['Low'] #, 'Low', 'Low', 'Low', 'Low', 'Low', 'Low','Low']
        source = ['RBI,ED,CBI,NIA'] #, 'UN,UK,EU,OFAC', 'UN,UK,EU,OFAC', 'UN,UK,EU,OFAC', 'UN,UK,EU,OFAC', 'UN,UK,EU,OFAC',
                  # 'UN,UK,EU,OFAC','UN,UK,EU,OFAC']
        output = []
        for i in range(0, len(name)):
            temp_dict = dict()
            temp_dict['name'] = name[i]
            temp_dict['sanaction_summary'] = sanaction_summary[i]
            temp_dict['sanaction_result'] = sanaction_result[i]
            temp_dict['sanaction_status'] = sanaction_status[i]
            temp_dict['source'] = source[i]
            output.append(temp_dict)
        # time.sleep(5)
        response = dict()
        response['output'] = output
        return response


@app.route('/missing_doc', methods=['GET'])
def missing_doc():
    attribute_list = [
        "Registered Office Address",
        "Legal Existence Of The Entity",
        "Nature of Business",
        "Industry Sectors with Higher Corruption Risk",
        "Country of Operation",
        "List of Senior Persons",
        "List of Legal Representatives"
    ]
    document_names = [
        "Economic Sanctions Due Diligence Questionnaire",
        "Board Minutes",
        "Articles of Association",
        "Memorandum of Association"
    ]
    response = dict()
    response['missing_attributes_list'] = attribute_list
    response['missing_docs_list'] = document_names
    return response

@app.route('/send_mail', methods=['POST'])
def send_mail():
    #This logic is for getting json
    request_dict = request.get_json()
    msg = Message(
        subject=request_dict['subject'],
        sender='ruleautomation8@gmail.com',
        recipients=request_dict['recipients']
    )
    msg.body = request_dict['body']
    # msg.html = request_dict['html']
    mail.send(msg)
    response = dict()
    response['message'] = 'Sent'
    return response

    # html_data = request.data  # Get the HTML data from the request
    #
    # # Parse the HTML content using BeautifulSoup
    # soup = BeautifulSoup(html_data, 'html.parser')
    #
    # # Extract values from the HTML using div IDs
    # subject = soup.find('div', {'id': 'subject'}).text
    # recipient = soup.find('div', {'id': 'Recipient'}).text
    # body = soup.find('div', {'id': 'body'}).text
    #
    # # Now you can use the subject, recipient, and body values in your Python code
    # print("Subject:", subject)
    # print("Recipient:", recipient)
    # print("Body:", body)
    # msg = Message(
    #           subject=subject,
    #           sender='ruleautomation8@gmail.com',
    #           recipients=[recipient]
    #       )
    # msg.body = body
    #
    # mail.send(msg)
    # response = dict()
    # response['message'] = 'Sent'
    # return response

##new method
@app.route("/mail_form_submission",methods=['POST'])
def mail_form_submission():
    req = request.get_json()
    if req['tab'] == 'missing_docs':
        mydb, cursor = dbConnection()
        for i in req['files']:
            encoded_text = i['fileContent']
            decoded_bytes = base64.b64decode(encoded_text)
            decoded_text = ''
            if i['fileType'].endswith('pdf'):
                pdf_file = BytesIO(decoded_bytes)
                decoded_text = extract_text(pdf_file)
                print(f"Type : {type(decoded_text)}decoded_text:{decoded_text}")
            if i['fileType'].startswith('text'):
                decoded_text = decoded_bytes.decode()
                # print(f"Type : {type(decoded_text)}decoded_text:{decoded_text}")
            if 'docx' in i['fileType'] or 'doc' in i['fileType']:
                docs_file = BytesIO(decoded_bytes)
                doc = docx.Document(docs_file)
                paragraphs = [p.text for p in doc.paragraphs]
                decoded_text = '\n'.join(paragraphs)
                print(f"Type : {type(decoded_text)}decoded_text:{decoded_text}")
            query = 'INSERT INTO fin_poc.entity_uploaded_file_storage (financial_year,entity_name,file_name, file_data) VALUES (%s,%s, %s,%s)'
            values = (req['finanical_year'], req['entity'], i['fileName'], decoded_text)
            cursor.execute(query, values)
            mydb.commit()

            print(f'{i["fileName"]}Details saved successfully in DB!!')
        response = dict()
        response['output'] = "File Details saved successfully in DB!!"
        return response

#Note: New Method
@app.route("/login_validation",methods=['POST'])
def login_validation():
    request_dict = request.get_json()
    if request_dict['tab'] == 'validation':
        username = request_dict["username"]
        password = request_dict["password"]
        mydb, cursor = dbConnection()
        sql_stmnt = "Select username,password,role from fin_poc.userlogin"
        cursor.execute(sql_stmnt)
        output = cursor.fetchall()
        db_username_li = [i for i,j,k in output]
        response = dict()
        if username in db_username_li:
            index = db_username_li.index(username)
            if password == output[index][1]:
                response['message'] = 'Login Validation Success !!'
                response['role'] = output[index][-1]
            else:
                response['message'] = 'Please check your Password !!'
        else:
            response['message'] = "Username doesn't exsists!!"
        print(response)
        return response
    if request_dict['tab'] == 'retrieve_role':
        user_val = request_dict["name"]
        # password = request_dict["password"]
        mydb, cursor = dbConnection()
        sql_stmnt = f"Select role from fin_poc.userlogin where username = '{user_val}'"
        print(sql_stmnt)
        cursor.execute(sql_stmnt)
        output = cursor.fetchone()
        print(output[0])
        response = dict()
        response['role'] = output[0]
        return response


#Note: New Method
@app.route("/media_pep_report_download",methods=['POST'])
def media_pep_report_download():
    req = request.get_json()
    if req['tab'] == "media_screen":
        if req['name'] == "Seamus Mulligan":
            filename = 'seamus_mulligan_scrapper_output.txt'
            return send_file(filename, as_attachment=True)
        if req['name'] == "Vijay Mallya":
            filename = 'Vijay_Mallaya_scraper_output.txt'
            return send_file(filename, as_attachment=True)
    if req['tab'] == "pep_screen":
        if req['name'] == "Amit Shah":
            filename = 'amit_shah_pep_output.txt'
            return send_file(filename, as_attachment=True)
        if req['name'] == "Kapil Sibal":
            filename = 'kapil_sibal_pep_output.txt'
            return send_file(filename, as_attachment=True)

    if req['tab'] == "sanaction_screen":
        if req['name'] == "Sadam Anwar":
            filename = 'sadam_anwar_sanaction_output.txt'
            return send_file(filename, as_attachment=True)

#Note: New Method
@app.route("/individual_cdd",methods=['GET'])
def individual_cdd():
    with open('sample6.png', 'rb') as image_file:
        image_data = image_file.read()
    aadhar_img = base64.b64encode(image_data).decode('utf-8')
    with open('pass_sample6.jpg', 'rb') as image_file:
        image_data = image_file.read()
    paasport_img = base64.b64encode(image_data).decode('utf-8')
    with open('dup.png', 'rb') as image_file:
        image_data = image_file.read()
    recent_aadhar_img = base64.b64encode(image_data).decode('utf-8')

    with open('pass_samp6_recent.png', 'rb') as image_file:
        image_data = image_file.read()
    recent_paasport_img = base64.b64encode(image_data).decode('utf-8')

    # Update the data for individual
    response = dict()
    response['ID Image'] = [aadhar_img,paasport_img]
    response['Proof Type'] = ['Aadhar','Passport - PE']
    response['Identification Number'] = ['938168750183','107185703']
    response['Name'] = ['Kamal Kumar','JODIE PIPPAS']
    response['Gender'] = ['Male','Female']
    response['CO'] = ['S/O: Lakhan Singh','-']
    response['Date of Birth'] = ['16/05/1989','17/01/1985']
    response['Address'] = ['Jamuna,Jamuna,Aligarh,Atrauli,Uttar Pradesh,202282','United Kingdom']
    response['Recent Image'] = [recent_aadhar_img,recent_paasport_img]
    response['result'] = ['Not Matched','Matched']
    response['score'] = ['21%', '93%']
    return response


@app.route("/cdd_image_comparison",methods=['GET'])
def cdd_image_comparison():
    with open('dup.png', 'rb') as image_file:
        image_data = image_file.read()
    aadhar_img = base64.b64encode(image_data).decode('utf-8')

    with open('pass_samp6_recent.png', 'rb') as image_file:
        image_data = image_file.read()
    paasport_img = base64.b64encode(image_data).decode('utf-8')

    response = dict()
    response['image'] = [aadhar_img,paasport_img]
    response['result'] = ['Not Matched','Matched']
    return response

@app.route("/edd_summary",methods=['GET'])
def edd_summary():
    li = ['The items which need confirmation from client are Source of Funds Entity,Entity Asset,Entity Liabilities',
          'The Beneficial owners are Kamal Kumar and Jodie Pippas',
          "Kamal Kumar's Aadhar document  " ]
    response = dict()
    response['output'] = li
    return response

@app.route("/download_entity_final_report",methods=['POST'])
def download_entity_final_report():
    req = request.get_json()
    name = req['entity']
    xls = pd.ExcelFile('entity_emergent_report_input_data_v2.xlsx')
    sheet_to_df_map = {}
    for sheet_name in xls.sheet_names:
        sheet_to_df_map[sheet_name] = xls.parse(sheet_name)

    output_filename = f"final_report_{name}.pdf"
    doc = SimpleDocTemplate(output_filename, pagesize=landscape(letter))

    elements = []

    # Add the title at the beginning of the document
    title = f"KYC Refresh Report for {name}(2023-24)"
    title_style = TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 20),
    ])
    title_table = Table([[title]], colWidths=[doc.width], rowHeights=[36])
    title_table.setStyle(title_style)
    elements.append(title_table)
    # elements.append(Spacer(1, 12))

    # Loop through each sheet and create a table for it
    for sheet_name, df in sheet_to_df_map.items():
        # Create a new page for each sheet
        elements.append(Spacer(1, 36))  # Add space for the sheet name

        # Add the sheet name in the middle of the page
        sheet_name_style = TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 15),
        ])
        sheet_name_table = Table([[sheet_name]], colWidths=[doc.width], rowHeights=[36])
        sheet_name_table.setStyle(sheet_name_style)
        elements.append(sheet_name_table)

        elements.append(Spacer(1, 12))  # Add space between sheets
        elements.append(Spacer(1, 12))
        data = [df.columns.values.tolist()] + df.values.tolist()
        table = Table(data)

        # Apply a TableStyle to format the table
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ])

        table.setStyle(table_style)
        # Create a KeepInFrame to allow horizontal scrolling
        table_in_frame = KeepInFrame(doc.width, doc.height - 120, [table])
        elements.append(table_in_frame)

    doc.build(elements)
    return send_file(output_filename, mimetype='application/pdf')


@app.route("/hierarchy_chart",methods=['POST'])
def hierarchy_chart():
    req = request.get_json()
    if req['tab'] == 'organization_hierarchy':
        domestic_childern = [
            {"name": "400 Professional LLC", "place": "Delaware"},
            {"name": "Cangene bioPharma, LLC", "place": "Maryland"},
            {"name": "Emergent Commercial Operations Frederick Inc.", "place": "Maryland"},
            {"name": "Emergent Biodefense Operations Lansing LLC", "place": "Delaware"},
            {"name": "Emergent Europe Inc.", "place": "Delaware"},
            {"name": "Emergent International Inc.", "place": "Delaware"},
            {"name": "Emergent Manufacturing Operations Baltimore LLC", "place": "Delaware"},
            {"name": "Emergent Product Development Gaithersburg Inc.", "place": "Delaware"},
            {"name": "Emergent Protective Products USA Inc.", "place": "Delaware"},
            {"name": "Emergent Virology LLC", "place": "Delaware"},
            {"name": "PaxVax Inc.", "place": "Delaware"},
            {"name": "Adapt Pharma Inc.", "place": "Delaware"}
        ]

        international_children = [
            {"name": "3579299 Manitoba Ltd.", "place": "Manitoba"},
            {"name": "Adapt Pharma Limited", "place": "Ireland"},
            {"name": "Adapt Pharma Canada Ltd.", "place": "British Columbia"},
            {"name": "Adapt Pharma Operations Limited", "place": "Ireland"},
            {"name": "Emergent Acquisition Limited", "place": "Ireland"},
            {"name": "Emergent BioSolutions Canada Inc. (f/k/a Cangene Corporation)", "place": "Ontario"},
            {"name": "Emergent BioSolutions Malaysia SDN. BHD.", "place": "Malaysia"},
            {"name": "Emergent BioSolutions Portugal, LDA", "place": "Portugal"},
            {"name": "Emergent Countermeasures International Ltd.", "place": "England"},
            {"name": "Emergent Global Health Foundation Limited", "place": "England"},
            {"name": "Emergent Italy S.r.l.", "place": "Italy"},
            {"name": "Emergent Netherlands B.V.", "place": "Netherlands"},
            {"name": "Emergent Product Development Germany GmbH", "place": "Germany"},
            {"name": "Emergent Product Development UK Limited", "place": "England"},
            {"name": "Emergent Sales and Marketing Australia Pty Ltd.", "place": "Australia"},
            {"name": "Emergent Sales and Marketing France S.A.S.", "place": "France"},
            {"name": "Emergent Sales and Marketing Germany GmbH", "place": "Germany"},
            {"name": "Emergent Sales and Marketing Singapore Pte. Ltd.", "place": "Singapore"},
            {"name": "Emergent BioSolutions UK LTD", "place": "England"},
            {"name": "PaxVax AUS Pty. Ltd.", "place": "Australia"},
            {"name": "PaxVax Bermuda Ltd.", "place": "Bermuda"},
            {"name": "PaxVax Berna GmbH", "place": "Switzerland"},
            {"name": "PaxVax Holding Company, Ltd.", "place": "Cayman Islands"},
            {"name": "PaxVax Spain, S.L.", "place": "Spain"}
        ]
        entity_children = [{'name':'Domestic','children':domestic_childern},{'name':'International','children':international_children}]
        response = dict()
        response['name'] = "Emergent Bio Solutions"
        response['children'] = entity_children
    if req['tab'] == 'ownership_hierarchy':
        non_employee_children=[
            {"name": "Dr. Sue Bailey", "percent_share": "<1%"},
            {"name": "Zsolt Harsanyi, Ph.D.", "percent_share": "<1%"},
            {"name": "Jerome Hauer, Ph.D.", "percent_share": "<1%"},
            {"name": "George Joulwan", "percent_share": "<1%"},
            {"name": "Seamus Mulligan", "percent_share": "1.1%"},
            {"name": "Ronald B. Richard", "percent_share": "<1%"},
            {"name": "Louis W. Sullivan, M.D.", "percent_share": "<1%"},
            {"name": "Kathryn Zoon, Ph.D.", "percent_share": "<1%"}
        ]
        executive_children = [
            {"name": "Fuad El-Hibri", "percent_share": "10.7%"},
            {"name": "Robert G. Kramer", "percent_share": "<1%"},
            {"name": "Richard S. Lindahl", "percent_share": "<1%"},
            {"name": "Adam Havey", "percent_share": "<1%"},
            {"name": "Atul Saran", "percent_share": "<1%"},
            {"name": "Other Executive Officers", "percent_share": "<1%"},
            {"name": "All executive officers and directors as a group (15 persons)", "percent_share": "13.1%"}
        ]
        stockholders_children = [
            {"name": "Vanguard Group", "percent_share": "8.9%"},
            {"name": "BlackRock, Inc.", "percent_share": "13.0%"},
            {"name": "Intervac, L.L.C.", "percent_share": "8.3%"}
        ]
        entity_children = [{'name':'Non-Employee Directors and Director Nominees','children':non_employee_children},{'name':'Named Executive Officers','children':executive_children},{'name':'5% or greater stockholders','children':stockholders_children}]
        response = dict()
        response['name'] = "Beneficiary Owners"
        response['children'] = entity_children
    return response

@app.route("/street_view",methods=['POST'])
def street_view():
    req = request.get_json()
    # api_key = os.getenv("GOOGLE_MAPS_API_KEY1")
    api_key = os.getenv("GOOGLE_MAPS_API_KEY")
    gmaps = googlemaps.Client(key=api_key)
    response = dict()
    try:
        # Geocode the location name to get its coordinates
        geocode_result = gmaps.geocode(req['location'])
        print("Geocode Result:", geocode_result)

        # location = geocode_result[0]['geometry']['location']
        # lat, lng = location['lat'], location['lng']
        lat = 51.4948334
        lng = -0.2761988
        # print(lat, lng)
        
        # Construct the Street View URL
        # street_view_url = f"https://www.google.com/maps/@?api=1&map_action=pano&viewpoint={lat},{lng}"
        response['latitude'] = lat
        response['longitude'] = lng
        response['api_key'] = api_key
        # response['street_view_url'] = street_view_url
        return response
    except Exception as e:
        print(f"Exception Type: {type(e).__name__}")
        print(f"Exception Message: {str(e)}")
        response['msg'] = str(e)
        return response


def read_image_extract_text(filepath):
    img = cv2.imread(filepath)
    # img = cv2.resize(img,(0,0),fx=0.15,fy=0.15)
    img = cv2.resize(img, (0, 0), fx=3, fy=3)
    gray_image = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # NOISE REMOVAL
    noise_removed_image = cv2.fastNlMeansDenoising(gray_image, None, 12, 15, 25)
    # INCREASE CONTRAST
    # sharp = filters.unsharp_mask(noise_removed_image, radius=2.25, amount=2.5,preserve_range=False)
    sharp = filters.unsharp_mask(noise_removed_image, radius=2.25, amount=2.5, preserve_range=False)

    sharp = (255 * sharp).clip(0, 255).astype(np.uint8)
    kernel = np.array([1])
    erosion = cv2.erode(sharp, kernel, iterations=10)
    img_dilation = cv2.dilate(erosion, kernel, iterations=10)

    text = pytesseract.image_to_string(img_dilation, lang='eng')
    text_output = open('output.txt', 'w', encoding='utf-8')
    text_output.write(text)
    text_output.close()
    file = open('output.txt', 'r', encoding='utf-8')
    text = file.read()
    text = ftfy.fix_text(text)
    text = ftfy.fix_encoding(text)
    #print(text)
    return text

def ocr_pan_read_data(text):
    name = None
    fname = None
    dob = None
    pan = None
    nameline = []
    dobline = []
    panline = []
    text0 = []
    text1 = []
    text2 = []
    lines = text.split('\n')
    for lin in lines:
        s = lin.strip()
        s = lin.replace('\n', '')
        s = s.rstrip()
        s = s.lstrip()
        text1.append(s)
    text1 = list(filter(None, text1))
    lineno = 0
    for wordline in text1:
        xx = wordline.split('\n')
        if ([w for w in xx if re.search('(INCOMETAXDEPARWENT|INCOME|TAX|GOW|GOVT|GOVERNMENT|OVERNMENT|VERNMENT|DEPARTMENT|EPARTMENT|PARTMENT|ARTMENT|INDIA|NDIA)$',
                w)]):
            text1 = list(text1)
            lineno = text1.index(wordline)
            break
    text0 = text1[lineno + 1:]

    for i in text0:
        if (re.search("[A-Z]{5}[0-9]{4}[A-Z]{1}", i)) and len(i) == 10:
            pan = i
        if (re.search(r"\d+/\d+/\d+", i)):
            dob = i
    try:
        # Cleaning first names
        name = text0[0]
        name = name.rstrip()
        name = name.lstrip()
        name = re.sub('[^a-zA-Z] +', ' ', name)
        # Cleaning Father's name
        fname = text0[1]
        fname = fname.rstrip()
        fname = fname.lstrip()
        fname = fname.replace("\"", "A")
        fname = re.sub('[^a-zA-Z] +', ' ', fname)
        # Cleaning DOB
        dob = dob.rstrip()
        dob = dob.lstrip()

    except:
        pass
    new_no = (len(pan) - 4) * 'x' + pan[-4:]
    id_no = ' '.join([new_no[i:i+4] for i in range(0, len(new_no), 4)])
    data = {}
    data['name'] = name
    data['father_name'] = fname
    data['dob'] = dob
    data['id_no'] = id_no
    #data['id_no'] = pan
    return data

def authentic_extract_aadhar_details(filepath):
    img = cv2.imread(filepath)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    code = decode(gray)
    try:
        qrData = code[0].data
        print(qrData, type(qrData))
    except:
        return "Please upoload clear image"
    xml_string = qrData.decode('utf-8')
    root = ET.fromstring(xml_string)
    attributes = root.attrib
    print(f'attributes:{attributes}')
    address_keys = []
    between_keys = False
    for key in attributes:
        if key == 'co':
            between_keys = True
        elif key == 'dob':
            break
        if between_keys:
            address_keys.append(key)
    # Concatenate the values of address keys
    address_values = ','.join(str(attributes.get(key, '')) for key in address_keys[1:])
    if attributes.get('gender') == 'M':
        attributes['gender'] = 'Male'
    elif attributes.get('gender') == 'F':
        attributes['gender'] = 'Female'
    aadhar_no = attributes.get('uid')
    new_no = (len(aadhar_no) - 4) * 'x' + aadhar_no[-4:]
    id_no = ' '.join([new_no[i:i + 4] for i in range(0, len(new_no), 4)])
    # details = {'id_no': id_no, 'name': attributes['name'], 'gender': attributes['gender'],
    #            "father_name": attributes['co'], 'dob': attributes['dob'], 'address':address_values}#, 'District':attributes['dist'], 'State':
    details = {'id_no': id_no, 'name': attributes['name'], 'dob': attributes['dob'],
               'address': address_values}  # , 'District':attributes['dist'], 'State':
    return details

def authentic_passport_extraction(filepath):
    mrz = read_mrz(filepath)
    mrz_data = mrz.to_dict()
    print(mrz_data)
    response = dict()
    nationalities_dict = {
        "USA": "United States of America",
        "CAN": "Canada",
        "GBR": "United Kingdom",
        "AUS": "Australia",
        "DEU": "Germany",
        "FRA": "France",
        "JPN": "Japan",
        "CHN": "China",
        "IND": "India",
        "BRA": "Brazil",
        "RUS": "Russia",
        "ZAF": "South Africa",
        "MEX": "Mexico",
        "ITA": "Italy",
        "ESP": "Spain",
        "ARG": "Argentina",
        "KOR": "South Korea",
        "SAU": "Saudi Arabia",
        "TUR": "Turkey",
        "EGY": "Egypt",
        "GRC": "Greece",
        "SWE": "Sweden",
        "NOR": "Norway",
        "DNK": "Denmark",
        "CHE": "Switzerland",
        "NLD": "Netherlands",
        "BEL": "Belgium",
        "IRL": "Ireland",
        "NZL": "New Zealand",
        "SGP": "Singapore",
        "MYS": "Malaysia",
        "THA": "Thailand",
        "VNM": "Vietnam",
        "IDN": "Indonesia",
        "PHL": "Philippines",
        "PAK": "Pakistan",
        "BGD": "Bangladesh",
        "LKA": "Sri Lanka",
        "NPL": "Nepal",
        "AFG": "Afghanistan",
        "IRQ": "Iraq",
        "IRN": "Iran",
        "ISR": "Israel",
        "JOR": "Jordan",
        "KWT": "Kuwait",
        "QAT": "Qatar",
        "ARE": "United Arab Emirates",
        "OMN": "Oman",
        "YEM": "Yemen",
        "BHR": "Bahrain"
    }
    response['Passport_Number'] = mrz_data['number']
    response['Name'] = mrz_data['names']
    response['Surname'] = mrz_data['surname']
    response['Nationality'] = nationalities_dict.get(mrz_data['nationality'])
    date_obj = datetime.strptime(mrz_data['date_of_birth'], "%y%m%d")
    formatted_date = date_obj.strftime("%d/%m/%Y")
    response['Date_of_birth'] = formatted_date
    response['Gender'] = 'Male'
    if mrz_data['sex'] == 'F':
        response['Gender'] = 'Female'
    response['Passport_Type'] = mrz_data['type']
    return response


@app.route('/individual_kyc', methods=['POST'])
def individual_kyc():
    req = request.get_json()

    if 'tab' in req:
        if req['tab'] == 'kyc_verification':
            return send_file('individual_kyc_details_output.xlsx', as_attachment=True)
        if req['tab'] == 'bar_graph_kyc_verification':
            return {"Total":8,"Compliant":3,"Non-Compliant":5}
        if req['tab'] == 'bar_graph_duplicate_acc':
            return {"Total":10,"Compliant":7,"Non-Compliant":3}
        if req['tab'] == 'download_report':
            return send_file('individual_kyc_compilance_report.xlsx', as_attachment=True)
        if req['tab'] == 'recent_image':
            image_path = "C:/Users/EK638MD/PycharmProjects/Fin_poc/Face_Compare/indivdual_cdd_samples/demo_samples/dup.png"
            with open(image_path, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read())
            #print(encoded_string)
            response = {"recent_image":encoded_string.decode("utf-8")}
            return response
    ##request body without home argument
    else:
        response = dict()
        for i in req:
            filepath = "output_image_conversion.jpg"
            base_64_data = req[i].split(',')[1]
            image_data = base64.b64decode(base_64_data)
            with open(filepath, "wb") as f:
                f.write(image_data)
            text = read_image_extract_text(filepath)
            print(f'text:{text}')
            if 'INCOME' in text:
                print('Inside Income')
                pan_extraction = ocr_pan_read_data(text)
                pan_extraction['cid'] = '2'
                pan_extraction['poi_db'] = 'PAN'
                pan_extraction['pov_db'] = 'Electricity Bill'
                pan_extraction['validity_check'] = 'Electricity Bill is not valid'
                pan_extraction['address'] = 'NA'
                pan_extraction['poi_match'] = 'Matched'
                pan_extraction['pov_match'] = 'NA'
                pan_extraction['compliant_status'] = 'Non-Compliant'
                pan_extraction['reason_non_complaint_area'] = 'POA Electricity Bill not Valid (Bill Date: 25/01/2022)'
                response[i] = pan_extraction
                #print(pan_extraction)
            elif 'Government of India' in text:
                print('Inside Aadhar')
                aadhar_extraction = authentic_extract_aadhar_details(filepath)
                #aadhar_extraction['type'] = 'Aadhar'
                aadhar_extraction['cid'] = '1'
                aadhar_extraction['poi_db'] = 'Aadhaar'
                aadhar_extraction['pov_db'] = 'Aadhaar'
                aadhar_extraction['validity_check'] = 'NA'
                aadhar_extraction['poi_match'] = 'Matched'
                aadhar_extraction['pov_match'] = 'Matched'
                aadhar_extraction['compliant_status'] = 'Compliant'
                aadhar_extraction['reason_non_complaint_area'] = 'NA'
                response[i] = aadhar_extraction
                #print(aadhar_extraction)
            elif 'REPUBLIC OF INDIA' in text or 'INDIAN' in text:
                print('Inside pass')
                #passport_extraction = authentic_passport_extraction(filepath)
                passport_extraction = dict()
                passport_extraction['cid'] = '4'
                passport_extraction['name'] = 'AMARJIT SINGH'
                passport_extraction['dob'] = '15/12/1994'
                # passport_extraction['id_no'] = 'K7758567'
                passport_extraction['id_no'] = 'xxxx 8567'
                passport_extraction['address'] = 'MIRANJANPUR ,AMRITSAR, PUNJAB'
                #passport_extraction['expiry_date'] = '24/03/2023'
                #passport_extraction['face_matching'] = 'Not Matched'
                passport_extraction['poi_db'] = 'Passport'
                passport_extraction['pov_db'] = 'Passport'
                passport_extraction['validity_check'] = 'Not Valid'
                passport_extraction['poi_match'] = 'NA'
                passport_extraction['pov_match'] = 'NA'
                passport_extraction['compliant_status'] = 'Non-Compliant'
                passport_extraction['reason_non_complaint_area'] = 'Passport Expired - (Exp. Date: 12/03/2023)'

                response[i] = passport_extraction
                #print(passport_extraction)
            elif "ELECTION COMMISSION OF INDIA" in text:
                voter_id_extraction = dict()
                # voter_id_extraction['type'] = 'Voter ID'
                voter_id_extraction['cid'] = '3'
                voter_id_extraction['name'] = 'Anindyasundar Mandal'
                #voter_id_extraction['id_no'] = 'LBT1381581'
                voter_id_extraction['id_no'] = 'xxxx xx15 81'
                voter_id_extraction['dob'] = 'NA'
                voter_id_extraction['address'] = 'Gopalmath Punabad,plot 35,Dungapur Burdwan,713217'
                voter_id_extraction['poi_db'] = 'Voter ID'
                voter_id_extraction['pov_db'] = 'Voter ID'
                voter_id_extraction['validity_check'] = 'NA'
                voter_id_extraction['poi_match'] = 'Matched'
                voter_id_extraction['pov_match'] = 'Not Matched'
                voter_id_extraction['compliant_status'] = 'Non-Compliant'
                voter_id_extraction['reason_non_complaint_area'] = 'Address not matched with Voter ID'
                response[i] = voter_id_extraction
                #print(voter_id_extraction)
            elif "MAHATMA GANDHI NATIONAL" in text:
                nrega_extraction = dict()
                nrega_extraction['cid'] = '5'
                nrega_extraction['name'] = 'SUNDER SINGH'
                # nrega_extraction['id_no'] = 'HR-09-002-071-001/119'
                nrega_extraction['id_no'] = 'xx-xx-xxx-xxx-xxx/119'
                nrega_extraction['address'] = "AHMADPUR,AHMADPUR,BALLABGARH"
                nrega_extraction['dob'] = "NA"
                nrega_extraction['poi_db'] = 'NREGA'
                nrega_extraction['pov_db'] = 'NREGA'
                nrega_extraction['validity_check'] = 'Not Valid'
                nrega_extraction['poi_match'] = 'NA'
                nrega_extraction['pov_match'] = 'NA'
                nrega_extraction['compliant_status'] = 'Non-Compliant'
                nrega_extraction['reason_non_complaint_area'] = 'Signature is missing in NREGA'

                response[i] = nrega_extraction
                #print(nrega_extraction)
        return response

@app.route('/reference_db', methods=['POST'])
def reference_db():
    # data = request.get_json()
    # excel_data = data
    # df = pd.DataFrame(excel_data[1:], columns=excel_data[0])
    # excel_filename = 'reference_db_output.xlsx'
    # df.to_excel(excel_filename, index=False)
    # return jsonify({'message': "Reference Database created successfully!! Please upload customer's KYC documents"})
    excel_blob = request.files['file'].read()
    df = pd.read_excel(io.BytesIO(excel_blob))
    print(f"before: {df['DOB'].to_list()}")
    #date_formats = pd.to_datetime(df['DOB'], errors='coerce').dt.strftime('%d/%m/%Y').unique()
    df['DOB'] = pd.to_datetime(df['DOB'], errors='coerce').dt.strftime('%d/%m/%Y')
    print(f"After: {df['DOB'].to_list()}")

    # for date_format in date_formats:
    #     df['DOB'] = pd.to_datetime(df['DOB'], errors='coerce', format=date_format).dt.strftime('%d/%m/%Y')
    json_response = [df.columns.tolist()] + df.values.tolist()
    print(json_response)
    return json_response



#@app.route("/send_mail_with_attachment_non_compliant",methods=['POST'])
@app.route("/email_RFI",methods=['POST'])
def send_emails_attach():
    non_compliant_user_dict = {
    "Cust ID" :["CUST001","CUST002","CUST003","CUST004","CUST005"],
    "CustomerID": ["Salaudin Ansari","Jasim Uddin", "Annindyasundar Mandal","John","Sunder Singh"],
    "CompliantStatus": ["Compliant", "Compliant", "Compliant","Non-Compliant","Compliant"],
    "EmailId":["souravadi1998@gmail.com", "mdjasimussin@gmail.com","mdjasimussin@gmail.com","biswasdipanjanaey24@gmail.com","mdjasimussin@gmail.com"] ,
    "reason":["", "","Document Invalid/Not Available",
              "W8 Form and Economic Sanctions Due Diligence Questionnaire are Invalid or Not Available", 
              "Document Invalid/Not Available"]}
    #non_compliant_user_dict = request.get_json()["data"]
    result = email_send(non_compliant_user_dict,"rfi").email_send()
    return jsonify(result)


@app.route("/send_mail_non_compliant",methods=['POST'])
@app.route("/email_EDD",methods=['POST'])
def send_emails():
    non_compliant_user_dict = {
    "Cust ID" :["CUST001","CUST002","CUST003","CUST004","CUST005"],
    "CustomerID": ["Salaudin Ansari","Jasim Uddin", "Annindyasundar Mandal","John","Sunder Singh"],
    "CompliantStatus": ["Compliant", "Compliant", "Compliant","Non-Compliant","Compliant"],
    "EmailId":["souravadi1998@gmail.com", "mdjasimussin@gmail.com","mdjasimussin@gmail.com","biswasdipanjanaey24@gmail.com","mdjasimussin@gmail.com"] ,
    "reason":["", "","Document Invalid/Not Available",
              "W8 Form and Economic Sanctions Due Diligence Questionnaire are Invalid or Not Available", 
              "Document Invalid/Not Available"]}
    #non_compliant_user_dict = request.get_json()["data"]
    result = email_send(non_compliant_user_dict,"edd").email_send()
    return jsonify(result)



@app.route('/download/<report_type>',methods=['POST'])
def download_reports(report_type):
    date = datetime.now().strftime('%Y%m%d')
    df_dict = request.get_json()["data"]
    df = pd.DataFrame(df_dict)
    if report_type == 'compliant':
        filtered_df = df[df['CompliantStatus'] == 'Compliant']
        filename = f'Compliant Customers - {date}.csv'
    elif report_type == 'non-compliant':
        filtered_df = df[df['CompliantStatus'] == 'Non-Compliant']
        filename = f'Non - Compliant Customers - {date}.csv'
    elif report_type == 'both':
        filtered_df = df
        filename = f'All Customers - {date}.csv'
    else:
        return "Invalid report type", 404

    res_data = filtered_df.to_dict('records')
    res_data = list(res_data)
    
    # # Convert DataFrame to CSV
    # buffer = io.StringIO()
    # filtered_df.to_csv(buffer, index=False)
    # buffer.seek(0)
    # csv_data = buffer.getvalue()

    # Create a response object with the CSV data
    # response = Response(csv_data, mimetype='text/csv')
    # response.headers.set("Content-Disposition", "attachment", filename=filename)
    # return response
    result = {"data":res_data,
              "filename":filename}
    return jsonify(result)


if __name__ == '__main__':
    app.run(port=5003,debug=True)



